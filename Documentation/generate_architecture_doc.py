from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Color Palette ──
DARK_BROWN = RGBColor(0x5C, 0x3A, 0x1E)
MEDIUM_BROWN = RGBColor(0x8B, 0x5E, 0x3C)
LIGHT_BROWN = RGBColor(0xA0, 0x72, 0x4D)
BEIGE_BG = RGBColor(0xFA, 0xF5, 0xEB)
BEIGE_ACCENT = RGBColor(0xE8, 0xDC, 0xC8)
BEIGE_ALT = RGBColor(0xF0, 0xE6, 0xD3)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x33, 0x33, 0x33)
BROWN_BORDER = RGBColor(0xC4, 0xA8, 0x82)

# ── Helpers ──
def set_cell_shading(cell, color_rgb):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_rgb}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_font(run, name='Calibri', size=11, color=BLACK, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold

def add_colored_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK_BROWN if level <= 1 else MEDIUM_BROWN
    return h

def set_paragraph_spacing(paragraph, before=0, after=6):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)

# ── Document Setup ──
doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = BLACK

# Set default paragraph spacing
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# ── Background colour for the document (section level) ──
section = doc.sections[0]
sect_pr = section._sectPr
bg = parse_xml(f'<w:background {nsdecls("w")} w:color="FAF5EB"/>')
sect_pr.append(bg)

# ── Title ──
title = doc.add_heading('II_Project — Architecture Documentation', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = DARK_BROWN
    run.font.size = Pt(26)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Bengos Restaurant Management System\nModule Architecture Specifications')
set_font(run, size=14, color=MEDIUM_BROWN)
set_paragraph_spacing(subtitle, before=0, after=18)

# ═══════════════════════════════════════════════════════════
# PART 1
# ═══════════════════════════════════════════════════════════
add_colored_heading(doc, '1. Architecture Patterns — How They Work', 1)

# ── 1.1 Code-Behind ──
add_colored_heading(doc, '1.1 Code-Behind Pattern', 2)

p = doc.add_paragraph()
r = p.add_run('Definition: ')
set_font(r, bold=True, color=DARK_BROWN)
r = p.add_run('The code-behind pattern is the default model for WPF and WinForms applications. The UI is defined in a markup/designer file (.xaml or .Designer.cs), and all logic — event handlers, data access, business rules — lives in a companion "code-behind" file (.xaml.cs or .cs).')
set_font(r)

add_colored_heading(doc, 'How it works:', 3)
for item in [
    'The XAML/Designer file declares UI elements (buttons, grids, textboxes).',
    'The code-behind file references these elements by name (e.g., txtUsername, btnLogin).',
    'Event handlers (Click, TextChanged, SelectionChanged) are wired to methods in the code-behind.',
    'All logic executes directly in these handlers — validation, database queries, calculations.',
    'There is no intermediate layer between the UI and the logic.'
]:
    doc.add_paragraph(item, style='List Bullet')

add_colored_heading(doc, 'Flow Diagram:', 3)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    '\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510    Click Event     \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510\n'
    '\u2502  XAML /  \u2502 \u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u25ba\u2502  Code-Behind.cs    \u2502\n'
    '\u2502 Designer \u2502\u25c4\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2502  (event handlers)  \u2502\n'
    '\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518   UI Update        \u2502  + SQL / File I/O  \u2502\n'
    '                                \u2502  + Business Logic   \u2502\n'
    '                                \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518'
)
set_font(r, 'Courier New', 9, MEDIUM_BROWN)

add_colored_heading(doc, 'Used in:', 3)
doc.add_paragraph('BengosRestaurantApp, Staff_Management, Billing&Payments, Invetory', style='List Bullet')

# ── 1.2 MVC ──
add_colored_heading(doc, '1.2 MVC (Model-View-Controller)', 2)

p = doc.add_paragraph()
r = p.add_run('Definition: ')
set_font(r, bold=True, color=DARK_BROWN)
r = p.add_run('MVC separates an application into three interconnected components: the Model (data + business rules), the View (UI), and the Controller (handles input, updates Model, selects View).')
set_font(r)

add_colored_heading(doc, 'How it works:', 3)
for item in [
    'Controller receives an HTTP request (e.g., GET /Menu/Category/Main).',
    'Controller interacts with the Model (via DbContext / repositories) to fetch or mutate data.',
    'Controller selects a View and passes the Model data to it (via ViewBag or a strongly-typed model).',
    'View renders HTML using Razor syntax, displaying the data from the Model.',
    'User interaction (click, form submit) sends a new HTTP request to the Controller — cycle repeats.'
]:
    doc.add_paragraph(item, style='List Bullet')

add_colored_heading(doc, 'Flow Diagram:', 3)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    '\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510   HTTP Request   \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510   Query    \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510\n'
    '\u2502  Browser \u2502\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u25ba\u2502 Controller \u2502\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u25ba\u2502 Model \u2502\n'
    '\u2502  (View)  \u2502\u25c4\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2502             \u2502\u25c4\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2502  +DB  \u2502\n'
    '\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518   HTML/JSON      \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518   Result   \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518'
)
set_font(r, 'Courier New', 9, MEDIUM_BROWN)

add_colored_heading(doc, 'Key characteristics:', 3)
for item in [
    'Separation of concerns: presentation (Razor Views), logic (Controllers), data (Models).',
    'Each Controller action is a concise method that typically does 3 things: get data, process, return view.',
    'Dependency Injection is used (e.g., RestaurantContext injected into MenuController).'
]:
    doc.add_paragraph(item, style='List Bullet')

add_colored_heading(doc, 'Used in:', 3)
doc.add_paragraph('DigitalClientMenu (BengosMenu) \u2014 the only module currently using a formal pattern.', style='List Bullet')

# ── 1.3 MVVM ──
add_colored_heading(doc, '1.3 MVVM (Model-View-ViewModel)', 2)

p = doc.add_paragraph()
r = p.add_run('Definition: ')
set_font(r, bold=True, color=DARK_BROWN)
r = p.add_run('MVVM is the recommended architecture for WPF applications. It separates the UI (View) from the presentation logic (ViewModel) and the data (Model), using data binding and commands as the communication mechanism.')
set_font(r)

add_colored_heading(doc, 'How it works:', 3)
for item in [
    'View (XAML) declares the UI and binds to properties and commands on the ViewModel.',
    'ViewModel exposes ObservableCollection properties, ICommand objects, and INotifyPropertyChanged events.',
    'View binds to ViewModel via DataContext (e.g., <Window DataContext="{Binding MainViewModel}">).',
    'When the user interacts with the UI, the binding engine automatically calls the corresponding command or updates the bound property.',
    'ViewModel calls Service/Repository methods to get/save data from the Model layer.',
    'Any change to a ViewModel property automatically updates the UI via data binding (two-way).'
]:
    doc.add_paragraph(item, style='List Bullet')

add_colored_heading(doc, 'Flow Diagram:', 3)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    '\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510  Data Binding   \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510  Calls   \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510  SQL/EF   \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2510\n'
    '\u2502  XAML    \u2502\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u25ba\u2502 ViewModel  \u2502\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u25ba\u2502  Service  \u2502\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u25ba\u2502 DB   \u2502\n'
    '\u2502  (View)  \u2502  INotifyPropCh  \u2502 (Commands) \u2502\u25c4\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2502 Repository\u2502\u25c4\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2502      \u2502\n'
    '\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518   + ICommand    \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518  Result  \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518  Data    \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2518'
)
set_font(r, 'Courier New', 9, MEDIUM_BROWN)

add_colored_heading(doc, 'Key characteristics:', 3)
for item in [
    'View knows about ViewModel; ViewModel knows about Model; Model knows nothing about the others.',
    'Two-way data binding: UI changes the ViewModel and ViewModel changes update the UI automatically.',
    'Testable: ViewModel logic can be unit-tested without the UI.',
    'Commands (ICommand) replace Click event handlers.'
]:
    doc.add_paragraph(item, style='List Bullet')

add_colored_heading(doc, 'Recommended for:', 3)
doc.add_paragraph('BengosRestaurantApp (WPF) \u2014 replace code-behind with ViewModels and data binding.', style='List Bullet')
doc.add_paragraph('inventoryWPF (WPF) \u2014 already has INotifyPropertyChanged on Product, just needs a ViewModel layer.', style='List Bullet')

# ── 1.4 MVP ──
add_colored_heading(doc, '1.4 MVP (Model-View-Presenter)', 2)

p = doc.add_paragraph()
r = p.add_run('Definition: ')
set_font(r, bold=True, color=DARK_BROWN)
r = p.add_run('MVP is the recommended architecture for WinForms applications. The View (Form) is passive and delegates all user actions to a Presenter, which contains the presentation logic and updates the View through an interface.')
set_font(r)

add_colored_heading(doc, 'How it works:', 3)
for item in [
    'View (Form) implements an IView interface with properties and events.',
    'Presenter receives a reference to the IView interface (not the concrete Form).',
    'View raises events (button click, selection changed) \u2192 Presenter handles them.',
    'Presenter calls Model / Service layer to get data.',
    'Presenter updates the View through the IView interface (sets properties like StaffName, ShiftList).',
    'View simply displays whatever the Presenter gives it \u2014 it contains no logic.'
]:
    doc.add_paragraph(item, style='List Bullet')

add_colored_heading(doc, 'Flow Diagram:', 3)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    '\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510   Interface   \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510   Calls   \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510\n'
    '\u2502  Form (View)   \u2502\u25c4\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u25ba\u2502 Presenter  \u2502\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u25ba\u2502 Model \u2502\n'
    '\u2502  (implements   \u2502    Events     \u2502            \u2502\u25c4\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2502  +DB  \u2502\n'
    '\u2502   IView)       \u2502   Properties  \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518  Result  \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518\n'
    '\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518'
)
set_font(r, 'Courier New', 9, MEDIUM_BROWN)

add_colored_heading(doc, 'Key characteristics:', 3)
for item in [
    'View is passive: it only displays data and forwards user actions.',
    'Presenter contains all presentation logic and is fully unit-testable.',
    'View-Presenter communication is through a defined interface (loose coupling).',
    'Better suited than MVVM for WinForms since WinForms lacks native data binding.'
]:
    doc.add_paragraph(item, style='List Bullet')

add_colored_heading(doc, 'Recommended for:', 3)
doc.add_paragraph('Staff_Management (WinForms) \u2014 extract logic from MainForm.cs into a Presenter.', style='List Bullet')
doc.add_paragraph('Billing&Payments (WinForms) \u2014 extract billing logic from Form1.cs into a Presenter.', style='List Bullet')

# ── Architecture Comparison Table ──
add_colored_heading(doc, '1.5 Architecture Comparison', 2)

table = doc.add_table(rows=5, cols=5)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set table borders to brown
tbl = table._tbl
tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    '  <w:top w:val="single" w:sz="6" w:space="0" w:color="C4A882"/>'
    '  <w:left w:val="single" w:sz="6" w:space="0" w:color="C4A882"/>'
    '  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="C4A882"/>'
    '  <w:right w:val="single" w:sz="6" w:space="0" w:color="C4A882"/>'
    '  <w:insideH w:val="single" w:sz="6" w:space="0" w:color="C4A882"/>'
    '  <w:insideV w:val="single" w:sz="6" w:space="0" w:color="C4A882"/>'
    '</w:tblBorders>'
)
tbl_pr.append(borders)

headers = ['Aspect', 'Code-Behind', 'MVC', 'MVVM', 'MVP']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    set_font(r, bold=True, size=10, color=WHITE)
    set_cell_shading(cell, '5C3A1E')

comp_data = [
    ['Platform', 'WPF / WinForms', 'ASP.NET Core (Web)', 'WPF', 'WinForms'],
    ['Separation', 'None', 'Controller\u2194Model\u2194View', 'View\u2194VM\u2194Model', 'View\u2194Presenter\u2194Model'],
    ['UI Updates', 'Direct (code)', 'HTTP Response (new page)', 'Data Binding (auto)', 'Via Interface (manual)'],
    ['Testability', 'Low', 'Medium', 'High', 'High'],
]
for row_idx, row_data in enumerate(comp_data):
    bg = BEIGE_ALT if row_idx % 2 == 0 else BEIGE_BG
    for col_idx, val in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(val)
        set_font(r, size=10)
        set_cell_shading(cell, bg)

doc.add_paragraph()  # spacer

# ═══════════════════════════════════════════════════════════
# PART 2 — Detail per Module
# ═══════════════════════════════════════════════════════════
add_colored_heading(doc, '2. Detailed Module Specifications', 1)

def add_module_table(doc, module_name, fields):
    add_colored_heading(doc, module_name, 2)
    table = doc.add_table(rows=len(fields), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Brown borders
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="6" w:space="0" w:color="C4A882"/>'
        '  <w:left w:val="single" w:sz="6" w:space="0" w:color="C4A882"/>'
        '  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="C4A882"/>'
        '  <w:right w:val="single" w:sz="6" w:space="0" w:color="C4A882"/>'
        '  <w:insideH w:val="single" w:sz="6" w:space="0" w:color="C4A882"/>'
        '  <w:insideV w:val="single" w:sz="6" w:space="0" w:color="C4A882"/>'
        '</w:tblBorders>'
    )
    tbl_pr.append(borders)

    # Set key column width
    for row in table.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(13)

    for i, (key, val) in enumerate(fields):
        # Key cell
        cell0 = table.rows[i].cells[0]
        cell0.text = ''
        p0 = cell0.paragraphs[0]
        r0 = p0.add_run(key)
        set_font(r0, bold=True, size=10, color=WHITE)
        set_cell_shading(cell0, '5C3A1E')

        # Value cell
        cell1 = table.rows[i].cells[1]
        cell1.text = ''
        p1 = cell1.paragraphs[0]
        r1 = p1.add_run(val)
        set_font(r1, size=10)
        bg = BEIGE_ALT if i % 2 == 0 else BEIGE_BG
        set_cell_shading(cell1, bg)

    doc.add_paragraph()

# ── Module 1 ──
add_module_table(doc, 'Module 1: BengosRestaurantApp (Main Hub)', [
    ('Project Path', 'BengosRestaurantApp/'),
    ('Technology', 'WPF (.NET 8.0-windows) \u2014 WinExe'),
    ('Current Architecture', 'Code-Behind (no formal pattern)'),
    ('Recommended Architecture', 'MVVM (Model-View-ViewModel)'),
    ('Description', 'Central dashboard application that launches all other restaurant management modules. Contains 7 windows: MainWindow (hub with navigation buttons), LoginWindow (user authentication from users.txt), StaffWindow (shift CRUD with role-based permissions), InventoryWindow (in-memory product CRUD), BillingWindow (order entry + discount), PaymentWindow (cash/card payment), MenuWindow (local menu browser with category pills).'),
    ('Data Access', 'Text files (users.txt, shifts.txt) via System.IO; in-memory ObservableCollection for orders and inventory.'),
    ('Current Problem', 'All logic \u2014 event handling, validation, file I/O, and business rules \u2014 lives in .xaml.cs code-behind files. No separation of concerns, hard to test, hard to maintain.'),
    ('MVVM Refactor', 'Extract ViewModels (MainViewModel, LoginViewModel, StaffViewModel, etc.) with ObservableProperties and ICommands. Move file I/O to a Repository layer. Bind XAML to ViewModel via DataContext.'),
])

# ── Module 2 ──
add_module_table(doc, 'Module 2: DigitalClientMenu (Web Digital Menu)', [
    ('Project Path', 'DigitalClientMenu/BengosMenu/'),
    ('Technology', 'ASP.NET Core MVC (.NET 10.0) \u2014 Web Application'),
    ('Current Architecture', 'MVC (Model-View-Controller) \u2705'),
    ('Recommended Architecture', 'MVC (keep \u2014 already correct)'),
    ('Description', 'Public-facing web application for customers to browse the restaurant menu. HomeController serves the landing page; MenuController serves category-filtered dish listings with ingredient details. Data is seeded on first run via Program.cs.'),
    ('Data Access', 'Entity Framework Core 10.0 + SQLite (restaurant.db). RestaurantContext with DbSet<Dish>, DbSet<Produs>, DbSet<DishIngredient>.'),
    ('Current Strength', 'Clean separation: Controllers handle HTTP, Models are entities, Data layer has DbContext, Views render HTML. DI for DbContext. This is the model module in the project.'),
    ('Note', 'No changes needed architecturally. Consider adding a Service layer between Controllers and DbContext for thicker business logic.'),
])

# ── Module 3 ──
add_module_table(doc, 'Module 3: Staff_Management (Shift Scheduling)', [
    ('Project Path', 'Staff_Management/hw1/'),
    ('Technology', 'WinForms (.NET 8.0-windows) \u2014 WinExe'),
    ('Current Architecture', 'Code-Behind (no formal pattern)'),
    ('Recommended Architecture', 'MVP (Model-View-Presenter)'),
    ('Description', 'Staff shift scheduling application. LoginForm authenticates users from users.txt (format: username,password,role). MainForm provides CRUD for shifts with role-based permissions: Admin can manage all staff shifts; regular staff can only manage their own. Shift data includes staff member, day, shift type (morning/evening/night), and overtime flag.'),
    ('Data Access', 'Text files (users.txt, shifts.txt) via System.IO.'),
    ('Current Problem', 'MainForm.cs is 583 lines \u2014 UI creation (manual), event handlers, file I/O, validation, and permissions logic are all in one class. WinForms designer file is not used (controls built in code).'),
    ('MVP Refactor', 'Create IShiftView interface (properties: StaffList, DayList, ShiftTypeList, SelectedShift, etc. + events: AddClicked, DeleteClicked). Create ShiftPresenter (handles events, calls ShiftRepository, updates View via interface). Keep Form thin \u2014 only UI setup and interface implementation.'),
])

# ── Module 4 ──
add_module_table(doc, 'Module 4: Billing & Payments', [
    ('Project Path', 'Billing&Payments/Homework/'),
    ('Technology', 'WinForms (.NET 8.0-windows) \u2014 WinExe'),
    ('Current Architecture', 'Code-Behind (no formal pattern)'),
    ('Recommended Architecture', 'MVP (Model-View-Presenter)'),
    ('Description', 'Order billing and payment processing. Form1 displays an order summary (hardcoded sample items: Cappuccino, Tiramisu, IceCream) with editable quantities, discount calculation popup, subtotal/discount/total display. Form2 handles payment \u2014 cash (with change calculation) or card.'),
    ('Data Access', 'In-memory (List<OrderItem>). No persistence.'),
    ('Current Problem', 'Order management, discount logic, and payment validation are all in Form1.cs and Form2.cs. Cannot unit-test billing rules without opening the Form.'),
    ('MVP Refactor', 'Create IBillingView interface and BillingPresenter. Move discount calculation and order total logic to Presenter. Create IPaymentView and PaymentPresenter for payment processing. Models (OrderItem) become proper domain classes.'),
])

# ── Module 5 ──
add_module_table(doc, 'Module 5: inventoryWPF (Inventory \u2014 WPF)', [
    ('Project Path', 'inventoryWPF/inventoryWPF/'),
    ('Technology', 'WPF (.NET 8.0-windows) \u2014 WinExe'),
    ('Current Architecture', 'Mixed: Code-Behind + INotifyPropertyChanged Model'),
    ('Recommended Architecture', 'MVVM (Model-View-ViewModel)'),
    ('Description', 'Product inventory management with CRUD operations against SQL LocalDB. Features: DataGrid display, Add/Edit/Delete products, category filtering (ComboBox), inventory alerts/notifications. Uses ObservableCollection<Product> bound to a DataGrid.'),
    ('Data Access', 'SQL LocalDB (Database_Products.mdf) via raw ADO.NET (SqlConnection, SqlCommand, SqlDataReader). System.Data.SqlClient 4.9.1.'),
    ('Current Strength', 'Product class implements INotifyPropertyChanged \u2014 data binding infrastructure is partially in place. ObservableCollection enables automatic UI updates on collection changes.'),
    ('Current Problem', 'SQL queries and CRUD logic are directly in MainWindow.xaml.cs event handlers (~240 lines). No separation between UI and data access. Connection string is hardcoded in code-behind.'),
    ('MVVM Refactor', 'Extract InventoryViewModel with ObservableCollection<Product>, ICommand for Add/Edit/Delete/Filter. Create IProductRepository interface with ProductRepository implementation (moves ADO.NET out of the View). Bind DataGrid and form fields to ViewModel.'),
])

# ── Module 6 ──
add_module_table(doc, 'Module 6: Invetory (Inventory \u2014 WinForms, Legacy)', [
    ('Project Path', 'Invetory/Invetory/'),
    ('Technology', 'WinForms (.NET 8.0-windows) \u2014 WinExe'),
    ('Current Architecture', 'Code-Behind (no formal pattern)'),
    ('Recommended Architecture', 'Deprecated \u2014 consolidate into inventoryWPF'),
    ('Description', 'Earlier/alternative inventory management application using WinForms. Single Form1 with standard WinForms CRUD against a database (connection string in App.config). Appears to be a prototype or alternate version of inventoryWPF.'),
    ('Data Access', 'Database via App.config connection string (likely LocalDB).'),
    ('Recommendation', 'This module duplicates inventoryWPF functionality. Recommended to archive or remove, consolidating all inventory management into the WPF version (module 5).'),
])

# ═══════════════════════════════════════════════════════════
# PART 3 — Module Dependency Map
# ═══════════════════════════════════════════════════════════
add_colored_heading(doc, '3. Module Dependency Map', 1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    '\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510\n'
    '\u2502                    BengosRestaurantApp                       \u2502\n'
    '\u2502                    (Main Hub \u2014 WPF)                          \u2502\n'
    '\u2502  \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510 \u2502\n'
    '\u2502  \u2502Inventory \u2502 \u2502  Staff   \u2502 \u2502 Billing  \u2502 \u2502  Menu (launch  \u2502 \u2502\n'
    '\u2502  \u2502 (in-app) \u2502 \u2502 (opens   \u2502 \u2502 (in-app) \u2502 \u2502  web browser)  \u2502 \u2502\n'
    '\u2502  \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518 \u2502 Staff_  \u2502 \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518 \u2502       \u2502        \u2502 \u2502\n'
    '\u2502               \u2502 Mgmt)   \u2502              \u2502  \u250c\u2500\u2500\u2500\u2500\u2514\u2500\u2500\u2500\u2500\u2500\u2510   \u2502 \u2502\n'
    '\u2502  Opens via    \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518   Opens via   \u2502  \u2502Digital   \u2502   \u2502 \u2502\n'
    '\u2502  button                    button       \u2502  \u2502ClientMenu\u2502   \u2502 \u2502\n'
    '\u2502               (WinForms)                \u2502  \u2502 (MVC)    \u2502   \u2502 \u2502\n'
    '\u2502                                          \u2502 \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518   \u2502 \u2502\n'
    '\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518\n'
    '\n'
    'Standalone modules (not launched from hub):\n'
    '  \u2022 inventoryWPF (WPF \u2014 separate .sln)\n'
    '  \u2022 Invetory (WinForms \u2014 legacy, standalone)'
)
set_font(r, 'Courier New', 9, MEDIUM_BROWN)

# ═══════════════════════════════════════════════════════════
# PART 4 — Master Table
# ═══════════════════════════════════════════════════════════
add_colored_heading(doc, '4. Master Summary Table', 1)

headers = ['Module', 'Technology', 'Current Arch', 'Recommended', 'Data Access', 'Description']
modules_data = [
    ['BengosRestaurantApp', 'WPF .NET 8', 'Code-Behind', 'MVVM', 'Text files, in-memory', 'Central hub launcher with 7 windows. Dashboard navigation to all sub-systems.'],
    ['DigitalClientMenu', 'ASP.NET Core MVC 10', 'MVC \u2705', 'MVC (keep)', 'EF Core + SQLite', 'Web digital menu. Customers browse dishes by category. Seed data.'],
    ['Staff_Management', 'WinForms .NET 8', 'Code-Behind', 'MVP', 'Text files\n(users.txt, shifts.txt)', 'Staff login and shift CRUD. Role-based permissions (Admin vs staff).'],
    ['Billing&Payments', 'WinForms .NET 8', 'Code-Behind', 'MVP', 'In-memory List', 'Order billing: qty edit, discount popup, cash/card payment.'],
    ['inventoryWPF', 'WPF .NET 8', 'Code-Behind + INPC', 'MVVM', 'SQL LocalDB (ADO.NET)', 'Product inventory CRUD with DataGrid and category filtering.'],
    ['Invetory', 'WinForms .NET 8', 'Code-Behind', 'Deprecated', 'DB (connection string)', 'Legacy WinForms inventory app. Duplicates module 5.'],
]

table = doc.add_table(rows=1 + len(modules_data), cols=6)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Brown borders
tbl = table._tbl
tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    '  <w:top w:val="single" w:sz="6" w:space="0" w:color="C4A882"/>'
    '  <w:left w:val="single" w:sz="6" w:space="0" w:color="C4A882"/>'
    '  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="C4A882"/>'
    '  <w:right w:val="single" w:sz="6" w:space="0" w:color="C4A882"/>'
    '  <w:insideH w:val="single" w:sz="6" w:space="0" w:color="C4A882"/>'
    '  <w:insideV w:val="single" w:sz="6" w:space="0" w:color="C4A882"/>'
    '</w:tblBorders>'
)
tbl_pr.append(borders)

num_rows = 1 + len(modules_data)
num_cols = 6
header_bg = '5C3A1E'

# Set column widths
widths = [Cm(3.5), Cm(3), Cm(3), Cm(3), Cm(3), Cm(5)]

for row_idx in range(num_rows):
    for col_idx in range(num_cols):
        cell = table.rows[row_idx].cells[col_idx]
        cell.width = widths[col_idx]
        cell.text = ''
        p = cell.paragraphs[0]

        if row_idx == 0:
            # Header row
            r = p.add_run(headers[col_idx])
            set_font(r, bold=True, size=9, color=WHITE)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, header_bg)
        else:
            # Data row
            val = modules_data[row_idx - 1][col_idx]
            r = p.add_run(val)
            set_font(r, size=9)
            bg = BEIGE_ALT if row_idx % 2 == 0 else BEIGE_BG
            set_cell_shading(cell, bg)

# ── Save ──
output_path = '/Users/user/Desktop/II_project/Documentation/II_Project_Architecture_Specification.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
